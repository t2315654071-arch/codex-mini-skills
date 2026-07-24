from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[1]
INPUT_DIR = REPO_ROOT / "examples" / "input"
OUTPUT_FILE = REPO_ROOT / "examples" / "output" / "提取结果.csv"
EXTRACTOR = (
    REPO_ROOT
    / "skills"
    / "document-extractor-lite"
    / "scripts"
    / "extract_lite.py"
)


def write_minimal_pdf(path: Path, text: str) -> None:
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(content)


def load_extractor():
    spec = importlib.util.spec_from_file_location("extract_lite", EXTRACTOR)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def main() -> None:
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)

    write_minimal_pdf(
        INPUT_DIR / "synthetic-invoice.pdf",
        "Synthetic Invoice INV-2026-001 Amount 128.50 Date 2026-07-24",
    )
    document = Document()
    document.add_heading("合成项目通知", level=1)
    document.add_paragraph("文档编号：SYN-2026-002")
    document.add_paragraph("日期：2026-07-24")
    document.add_paragraph("本文件完全由程序生成，不包含任何客户数据。")
    document.save(INPUT_DIR / "合成项目通知.docx")

    result = load_extractor().extract_directory(INPUT_DIR, OUTPUT_FILE)
    print(result)


if __name__ == "__main__":
    main()
