from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

VERSION = "1.0.0"
MAX_FILES = 10
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
CSV_FIELDS = [
    "文件名",
    "格式",
    "标题",
    "字符数",
    "页数或段落数",
    "文本预览",
    "状态",
    "错误",
]


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def extract_pdf(path: Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError("PDF已加密，免费版不会尝试破解")
    pages = [page.extract_text() or "" for page in reader.pages]
    text = normalize_text("\n".join(pages))
    if not text:
        raise ValueError("PDF没有可提取文字，可能是扫描件或空白文档")
    return text, len(reader.pages)


def extract_docx(path: Path) -> tuple[str, int]:
    document = Document(str(path))
    blocks: list[str] = []
    paragraph_count = 0
    for paragraph in document.paragraphs:
        value = normalize_text(paragraph.text)
        if value:
            blocks.append(value)
            paragraph_count += 1
    for table in document.tables:
        for row in table.rows:
            value = normalize_text(" | ".join(cell.text for cell in row.cells))
            if value:
                blocks.append(value)
                paragraph_count += 1
    text = normalize_text("\n".join(blocks))
    if not text:
        raise ValueError("DOCX没有可提取文字")
    return text, paragraph_count


def build_row(path: Path) -> dict[str, Any]:
    row: dict[str, Any] = {
        "文件名": path.name,
        "格式": path.suffix.lower().lstrip("."),
        "标题": "",
        "字符数": 0,
        "页数或段落数": 0,
        "文本预览": "",
        "状态": "失败",
        "错误": "",
    }
    try:
        if path.suffix.lower() == ".pdf":
            text, count = extract_pdf(path)
        else:
            text, count = extract_docx(path)
        title = text[:80]
        row.update(
            {
                "标题": title,
                "字符数": len(text),
                "页数或段落数": count,
                "文本预览": text[:300],
                "状态": "成功",
            }
        )
    except Exception as exc:
        row["错误"] = f"{type(exc).__name__}: {exc}"
    return row


def extract_directory(input_dir: Path, output: Path) -> dict[str, Any]:
    input_dir = input_dir.resolve()
    if not input_dir.is_dir():
        raise ValueError(f"输入目录不存在: {input_dir}")
    files = sorted(
        (
            path
            for path in input_dir.iterdir()
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ),
        key=lambda path: path.name.casefold(),
    )
    if not files:
        raise ValueError("输入目录中没有受支持的PDF或DOCX")
    if len(files) > MAX_FILES:
        raise ValueError(f"免费版单次最多处理{MAX_FILES}个文件，当前为{len(files)}个")

    rows = [build_row(path) for path in files]
    output = output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=CSV_FIELDS)
        writer.writeheader()
        writer.writerows(rows)
    success = sum(row["状态"] == "成功" for row in rows)
    return {
        "status": "ok",
        "version": VERSION,
        "processed": len(rows),
        "success": success,
        "failed": len(rows) - success,
        "output": str(output),
    }


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description="文档提取免费版：PDF/DOCX转固定CSV。")
    parser.add_argument("--input-dir", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--version", action="store_true")
    args = parser.parse_args()
    if args.version:
        print(VERSION)
        return 0
    if args.input_dir is None or args.output is None:
        parser.error("--input-dir和--output为必填参数")
    try:
        print(
            json.dumps(
                extract_directory(args.input_dir, args.output),
                ensure_ascii=False,
                indent=2,
            )
        )
        return 0
    except Exception as exc:
        print(
            json.dumps(
                {"status": "error", "error": f"{type(exc).__name__}: {exc}"},
                ensure_ascii=False,
            ),
            file=sys.stderr,
        )
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
